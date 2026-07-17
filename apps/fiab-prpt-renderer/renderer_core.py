"""Pure render core for the CSA Loom paginated-report renderer.

Extracted verbatim from azure-functions/paginated-report-renderer/function_app.py
(the Azure Functions host cannot run in estates whose policy forces
publicNetworkAccess=Disabled on storage — a Y1 consumption plan has no VNet
integration, so the host cannot reach its own AzureWebJobsStorage and never
starts; diagnosed live 2026-07-15). The ACA host in app.py serves the same
/api/render + /api/health contract in-VNet.
"""

from __future__ import annotations

import io
import logging
from typing import Any

logger = logging.getLogger("loom.paginated_report_renderer")

# --- ReportLab page geometry -------------------------------------------------
_PAGE_SIZES = {"A4": "A4", "Letter": "LETTER", "Legal": "LEGAL"}


# ---------------------------------------------------------------------------
# Expression evaluation
# ---------------------------------------------------------------------------

_AGGS = ("Sum", "Count", "Avg", "Max", "Min")


def _parse_expr(expr: str) -> tuple[str, str]:
    """Return (agg, field) for a cell expression. agg='' for a plain field."""
    if not expr:
        return ("", "")
    expr = expr.strip()
    if expr.startswith("=") and expr.endswith(")") and "(" in expr:
        head, rest = expr[1:].split("(", 1)
        inner = rest[:-1]
        agg = head.strip()
        field = _field_name(inner)
        if agg in _AGGS:
            return (agg, field)
    return ("", _field_name(expr))


def _field_name(token: str) -> str:
    token = token.strip()
    if token.startswith("Fields!") and token.endswith(".Value"):
        return token[len("Fields!"):-len(".Value")]
    return token


def _to_number(v: Any) -> float | None:
    try:
        if v is None or isinstance(v, bool):
            return None
        return float(v)
    except (TypeError, ValueError):
        return None


def _aggregate(agg: str, field: str, rows: list[dict]) -> Any:
    vals = [r.get(field) for r in rows]
    if agg == "Count":
        return sum(1 for v in vals if v is not None)
    nums = [n for n in (_to_number(v) for v in vals) if n is not None]
    if not nums:
        return ""
    if agg == "Sum":
        return round(sum(nums), 4)
    if agg == "Avg":
        return round(sum(nums) / len(nums), 4)
    if agg == "Max":
        return max(nums)
    if agg == "Min":
        return min(nums)
    return ""


def _cell_value(expr: str, row: dict) -> Any:
    agg, field = _parse_expr(expr)
    if agg:
        # Aggregates render in the totals row, not per detail row.
        return ""
    val = row.get(field, "")
    return "" if val is None else val


def _has_aggregate(tablix: dict) -> bool:
    for cell_row in tablix.get("cells", []):
        for cell in cell_row:
            agg, _ = _parse_expr((cell or {}).get("expression", ""))
            if agg:
                return True
    return False


def _tablix_rows(tablix: dict, sample_rows: list[dict]) -> tuple[list[str], list[list[str]], list[str] | None]:
    """Return (header, detail_rows, totals_row|None) as display strings."""
    columns: list[str] = tablix.get("columns", [])
    header_row: list[str] = tablix.get("headerRow", []) or columns
    cells = tablix.get("cells", [[]])
    detail_template = cells[0] if cells else []

    def expr_for(ci: int) -> str:
        if ci < len(detail_template) and detail_template[ci]:
            return detail_template[ci].get("expression", "")
        col = columns[ci] if ci < len(columns) else ""
        return f"Fields!{col}.Value"

    detail: list[list[str]] = []
    for row in sample_rows:
        detail.append([str(_cell_value(expr_for(ci), row)) for ci in range(len(columns))])

    totals: list[str] | None = None
    if _has_aggregate(tablix):
        totals = []
        for ci in range(len(columns)):
            agg, field = _parse_expr(expr_for(ci))
            totals.append(str(_aggregate(agg, field, sample_rows)) if agg else "")
    header = [str(h) for h in (header_row[: len(columns)] + columns[len(header_row):])][: len(columns)]
    if not header:
        header = columns
    return header, detail, totals


def _datasets_by_id(definition: dict) -> dict[str, dict]:
    return {d.get("id"): d for d in definition.get("datasets", [])}


# ---------------------------------------------------------------------------
# PDF (ReportLab)
# ---------------------------------------------------------------------------

def _render_pdf(definition: dict) -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, LETTER, LEGAL, landscape
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak

    size_map = {"A4": A4, "Letter": LETTER, "Legal": LEGAL}
    page = size_map.get(definition.get("pageSize", "Letter"), LETTER)
    if definition.get("pageOrientation") == "Landscape":
        page = landscape(page)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=page, title=definition.get("name", "report"))
    styles = getSampleStyleSheet()
    story: list[Any] = [Paragraph(definition.get("name", "Paginated report"), styles["Title"])]
    if definition.get("description"):
        story.append(Paragraph(str(definition["description"]), styles["Normal"]))
    story.append(Spacer(1, 12))

    dsets = _datasets_by_id(definition)
    tablixes = definition.get("tablixes", [])
    for ti, tablix in enumerate(tablixes):
        ds = dsets.get(tablix.get("datasetId"), {})
        sample_rows = ds.get("sampleRows", []) or []
        header, detail, totals = _tablix_rows(tablix, sample_rows)
        story.append(Paragraph(tablix.get("name", "Table"), styles["Heading2"]))

        data: list[list[str]] = []
        if tablix.get("showColumnHeaders", True):
            data.append(header)
        data.extend(detail if detail else [["(no rows)"] + [""] * (len(header) - 1)])
        if totals:
            data.append(totals)

        tbl = Table(data, repeatRows=1 if tablix.get("showColumnHeaders", True) else 0)
        style = [
            ("FONTSIZE", (0, 0), (-1, -1), 9),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#B0B0B0")),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 4),
            ("RIGHTPADDING", (0, 0), (-1, -1), 4),
        ]
        if tablix.get("showColumnHeaders", True):
            style += [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2B3A67")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ]
        if totals:
            style += [("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
                      ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#EFEFEF"))]
        tbl.setStyle(TableStyle(style))
        story.append(tbl)
        story.append(Spacer(1, 16))
        if tablix.get("pageBreak") and ti < len(tablixes) - 1:
            story.append(PageBreak())

    if not tablixes:
        story.append(Paragraph("This report has no tablix.", styles["Italic"]))

    doc.build(story)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# XLSX (openpyxl)
# ---------------------------------------------------------------------------

def _render_xlsx(definition: dict) -> bytes:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    wb = Workbook()
    wb.remove(wb.active)  # drop default sheet; one sheet per tablix
    dsets = _datasets_by_id(definition)
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="2B3A67", end_color="2B3A67", fill_type="solid")
    totals_font = Font(bold=True)

    used = set()
    for idx, tablix in enumerate(definition.get("tablixes", [])):
        raw = (tablix.get("name") or f"Table{idx+1}")[:28]
        title = raw or f"Table{idx+1}"
        n = title
        c = 1
        while n.lower() in used:
            c += 1
            n = f"{title[:25]}_{c}"
        used.add(n.lower())
        ws = wb.create_sheet(n)

        ds = dsets.get(tablix.get("datasetId"), {})
        sample_rows = ds.get("sampleRows", []) or []
        header, detail, totals = _tablix_rows(tablix, sample_rows)

        r = 1
        if tablix.get("showColumnHeaders", True):
            for ci, h in enumerate(header, start=1):
                cell = ws.cell(row=r, column=ci, value=h)
                cell.font = header_font
                cell.fill = header_fill
            r += 1
        for drow in detail:
            for ci, v in enumerate(drow, start=1):
                ws.cell(row=r, column=ci, value=v)
            r += 1
        if totals:
            for ci, v in enumerate(totals, start=1):
                cell = ws.cell(row=r, column=ci, value=v)
                cell.font = totals_font

        # Auto-ish column widths.
        for ci in range(1, max(1, len(header)) + 1):
            longest = len(str(header[ci - 1])) if ci - 1 < len(header) else 8
            for drow in detail:
                if ci - 1 < len(drow):
                    longest = max(longest, len(str(drow[ci - 1])))
            ws.column_dimensions[get_column_letter(ci)].width = min(60, max(10, longest + 2))

    if not definition.get("tablixes"):
        ws = wb.create_sheet("Report")
        ws["A1"] = "This report has no tablix."

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DOCX (python-docx)
# ---------------------------------------------------------------------------

def _render_docx(definition: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt

    doc = Document()
    doc.add_heading(definition.get("name", "Paginated report"), level=0)
    if definition.get("description"):
        doc.add_paragraph(str(definition["description"]))

    dsets = _datasets_by_id(definition)
    tablixes = definition.get("tablixes", [])
    for ti, tablix in enumerate(tablixes):
        ds = dsets.get(tablix.get("datasetId"), {})
        sample_rows = ds.get("sampleRows", []) or []
        header, detail, totals = _tablix_rows(tablix, sample_rows)
        doc.add_heading(tablix.get("name", "Table"), level=2)

        ncols = max(1, len(header))
        body_rows = detail if detail else [["(no rows)"] + [""] * (ncols - 1)]
        total_rows = (1 if tablix.get("showColumnHeaders", True) else 0) + len(body_rows) + (1 if totals else 0)
        table = doc.add_table(rows=total_rows, cols=ncols)
        table.style = "Light Grid Accent 1"

        ri = 0
        if tablix.get("showColumnHeaders", True):
            for ci in range(ncols):
                cell = table.cell(ri, ci)
                cell.text = str(header[ci]) if ci < len(header) else ""
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True
                        run.font.size = Pt(10)
            ri += 1
        for drow in body_rows:
            for ci in range(ncols):
                table.cell(ri, ci).text = str(drow[ci]) if ci < len(drow) else ""
            ri += 1
        if totals:
            for ci in range(ncols):
                cell = table.cell(ri, ci)
                cell.text = str(totals[ci]) if ci < len(totals) else ""
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.bold = True

        if tablix.get("pageBreak") and ti < len(tablixes) - 1:
            doc.add_page_break()

    if not tablixes:
        doc.add_paragraph("This report has no tablix.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Render entrypoint
# ---------------------------------------------------------------------------

_RENDERERS = {"pdf": _render_pdf, "xlsx": _render_xlsx, "docx": _render_docx}
_MIME = {
    "pdf": "application/pdf",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


