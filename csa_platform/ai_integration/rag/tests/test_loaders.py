"""Tests for the PDF / DOCX loader modules (CSA-0097).

The ``pypdf`` + ``python-docx`` integration paths are exercised against
small fixtures committed under ``tests/fixtures/``.  The Azure Document
Intelligence path is gated on ``RAG_DOC_INTELLIGENCE_ENDPOINT`` and is
covered with a fully mocked client so no Azure call is made.
"""

from __future__ import annotations

from pathlib import Path
from typing import Any
from unittest.mock import MagicMock, patch

import pytest

from csa_platform.ai_integration.rag import loaders

FIXTURES = Path(__file__).parent / "fixtures"


class TestLoadPdfPypdf:
    def test_extracts_two_pages_from_sample(self) -> None:
        segments = loaders._load_pdf_pypdf(FIXTURES / "sample.pdf")
        assert len(segments) == 2
        assert segments[0][1] == "Page 1"
        assert segments[1][1] == "Page 2"
        assert "Sample PDF" in segments[0][0]
        assert "Second Page" in segments[1][0]

    def test_load_pdf_falls_through_to_pypdf_without_endpoint(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.delenv("RAG_DOC_INTELLIGENCE_ENDPOINT", raising=False)
        segments = loaders.load_pdf(FIXTURES / "sample.pdf")
        assert len(segments) == 2
        assert segments[0][1] == "Page 1"

    def test_load_pdf_uses_doc_intelligence_when_endpoint_set(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.setenv(
            "RAG_DOC_INTELLIGENCE_ENDPOINT", "https://fake-docintel.example"
        )

        def _fake_di(path: Path, endpoint: str) -> list[tuple[str, str | None]]:
            assert endpoint == "https://fake-docintel.example"
            return [("DI page one content", "Page 1")]

        monkeypatch.setattr(loaders, "_load_pdf_document_intelligence", _fake_di)
        segments = loaders.load_pdf(FIXTURES / "sample.pdf")
        assert segments == [("DI page one content", "Page 1")]

    def test_load_pdf_falls_back_when_doc_intelligence_raises(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        monkeypatch.setenv(
            "RAG_DOC_INTELLIGENCE_ENDPOINT", "https://fake-docintel.example"
        )

        def _boom(path: Path, endpoint: str) -> list[tuple[str, str | None]]:
            raise RuntimeError("docintel exploded")

        monkeypatch.setattr(loaders, "_load_pdf_document_intelligence", _boom)
        segments = loaders.load_pdf(FIXTURES / "sample.pdf")
        # pypdf fallback always produces real pages.
        assert len(segments) == 2
        assert segments[0][1] == "Page 1"

    def test_doc_intelligence_helper_with_mocked_client(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        """Mock the Document Intelligence SDK surface end-to-end."""
        line_a = MagicMock()
        line_a.content = "Hello world"
        line_b = MagicMock()
        line_b.content = "Second line"
        page = MagicMock()
        page.page_number = 1
        page.lines = [line_a, line_b]

        result = MagicMock()
        result.pages = [page]
        poller = MagicMock()
        poller.result.return_value = result
        client = MagicMock()
        client.begin_analyze_document.return_value = poller

        fake_di_module = MagicMock()
        fake_di_module.DocumentIntelligenceClient = MagicMock(return_value=client)
        fake_di_models = MagicMock()
        fake_di_models.AnalyzeResult = MagicMock

        fake_core_creds = MagicMock()
        fake_core_creds.AzureKeyCredential = MagicMock(
            return_value=MagicMock(name="AzureKeyCredential")
        )

        monkeypatch.setenv("RAG_DOC_INTELLIGENCE_KEY", "fake-key")
        with patch.dict(
            "sys.modules",
            {
                "azure.ai.documentintelligence": fake_di_module,
                "azure.ai.documentintelligence.models": fake_di_models,
                "azure.core.credentials": fake_core_creds,
            },
        ):
            segments = loaders._load_pdf_document_intelligence(
                FIXTURES / "sample.pdf", endpoint="https://fake-di.example"
            )
        assert segments == [("Hello world\nSecond line", "Page 1")]


class TestLoadDocx:
    def test_extracts_heading_sections(self) -> None:
        segments = loaders.load_docx(FIXTURES / "sample.docx")
        # Introduction section + Setup section = 2 segments.
        anchors = [anchor for _, anchor in segments]
        assert "Heading: Introduction" in anchors
        assert "Heading: Setup" in anchors

    def test_segments_contain_body_text(self) -> None:
        segments = loaders.load_docx(FIXTURES / "sample.docx")
        # Find the setup segment and confirm its body is attached.
        setup_bodies = [text for text, anchor in segments if anchor == "Heading: Setup"]
        assert setup_bodies
        assert "Setup instructions" in setup_bodies[0]

    def test_docx_without_headings_yields_single_segment(
        self, tmp_path: Path
    ) -> None:
        import docx

        doc = docx.Document()
        doc.add_paragraph("Just a body paragraph without any headings at all.")
        doc.add_paragraph("Another body paragraph to make the segment non-trivial.")
        path = tmp_path / "no-headings.docx"
        doc.save(str(path))

        segments = loaders.load_docx(path)
        assert len(segments) == 1
        text, anchor = segments[0]
        assert anchor is None
        assert "body paragraph" in text


class TestPypdfMissingDep:
    def test_pdf_loader_raises_runtime_error_without_pypdf(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import builtins

        real_import = builtins.__import__

        def _fail(name: str, *args: Any, **kwargs: Any) -> Any:
            if name == "pypdf":
                raise ImportError("pypdf not installed")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", _fail)
        with pytest.raises(RuntimeError, match="pypdf"):
            loaders._load_pdf_pypdf(FIXTURES / "sample.pdf")


class TestDocxMissingDep:
    def test_docx_loader_raises_runtime_error_without_python_docx(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        import builtins

        real_import = builtins.__import__

        def _fail(name: str, *args: Any, **kwargs: Any) -> Any:
            if name == "docx":
                raise ImportError("python-docx not installed")
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", _fail)
        with pytest.raises(RuntimeError, match="python-docx"):
            loaders.load_docx(FIXTURES / "sample.docx")
