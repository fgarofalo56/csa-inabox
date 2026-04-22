"""Multi-format document loader for GraphRAG knowledge store.

Supports: PDF (via Azure Document Intelligence), CSV, DOCX, Markdown,
and GitHub repository scanning. Uploads processed documents to Azure
Blob Storage for GraphRAG indexing.
"""

from __future__ import annotations

import glob
import json
import logging
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

from azure.identity import DefaultAzureCredential
from azure.storage.blob import BlobServiceClient

logger = logging.getLogger(__name__)


@dataclass
class Document:
    """A processed document ready for GraphRAG indexing."""

    title: str
    content: str
    source: str
    metadata: dict[str, Any] = field(default_factory=dict)

    def to_text(self) -> str:
        """Format as text file content for GraphRAG input."""
        header = f"# {self.title}\n\nSource: {self.source}\n\n"
        return header + self.content


class DocumentLoader:
    """Load documents from multiple formats for GraphRAG knowledge store.

    Supports PDF, CSV, DOCX, Markdown, and GitHub repository scanning.
    Outputs text files to a staging directory or uploads directly to
    Azure Blob Storage for GraphRAG indexing.

    Example:
        loader = DocumentLoader()
        docs = loader.load_directory("./documents", formats=["pdf", "md", "docx"])
        loader.upload_to_blob(docs, container_name="graphrag-input")
    """

    def __init__(
        self,
        document_intelligence_endpoint: str | None = None,
        credential: Any | None = None,
    ) -> None:
        self._di_endpoint = document_intelligence_endpoint or os.getenv(
            "AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT"
        )
        self._credential = credential or DefaultAzureCredential()

    def load_directory(
        self,
        directory: str | Path,
        formats: list[str] | None = None,
        recursive: bool = True,
    ) -> list[Document]:
        """Load all supported documents from a directory.

        Args:
            directory: Path to directory containing documents.
            formats: List of formats to load. Default: all supported.
                Supported: pdf, csv, docx, md, txt, json.
            recursive: Whether to search subdirectories.

        Returns:
            List of processed Document objects.
        """
        directory = Path(directory)
        if not directory.is_dir():
            raise FileNotFoundError(f"Directory not found: {directory}")

        format_map = {
            "pdf": "*.pdf",
            "csv": "*.csv",
            "docx": "*.docx",
            "md": "*.md",
            "txt": "*.txt",
            "json": "*.json",
        }
        formats = formats or list(format_map.keys())
        documents: list[Document] = []

        for fmt in formats:
            pattern = format_map.get(fmt)
            if not pattern:
                logger.warning("Unsupported format: %s", fmt)
                continue

            glob_pattern = f"**/{pattern}" if recursive else pattern
            files = list(directory.glob(glob_pattern))
            logger.info("Found %d %s files in %s", len(files), fmt, directory)

            for file_path in files:
                try:
                    docs = self._load_file(file_path, fmt)
                    documents.extend(docs)
                except Exception:
                    logger.exception("Failed to load %s", file_path)

        logger.info("Loaded %d documents total", len(documents))
        return documents

    def load_pdfs(self, directory: str | Path) -> list[Document]:
        """Load PDF files using Azure Document Intelligence.

        Requires AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT environment variable
        or endpoint passed to constructor.
        """
        return self.load_directory(directory, formats=["pdf"])

    def load_csvs(
        self,
        directory: str | Path,
        text_columns: list[str] | None = None,
        title_column: str | None = None,
        max_rows_per_doc: int = 100,
    ) -> list[Document]:
        """Load CSV files, converting rows to document text.

        Args:
            directory: Path to CSV files.
            text_columns: Columns to include in document text.
                If None, all columns are included.
            title_column: Column to use as document title.
            max_rows_per_doc: Maximum rows per output document.
                Large CSVs are split into multiple documents.
        """
        import csv

        directory = Path(directory)
        documents: list[Document] = []

        for csv_path in directory.glob("**/*.csv"):
            try:
                with open(csv_path, newline="", encoding="utf-8-sig") as f:
                    reader = csv.DictReader(f)
                    rows = list(reader)

                if not rows:
                    continue

                columns = text_columns or list(rows[0].keys())

                # Split into chunks
                for i in range(0, len(rows), max_rows_per_doc):
                    chunk = rows[i : i + max_rows_per_doc]
                    chunk_num = i // max_rows_per_doc + 1
                    total_chunks = (len(rows) + max_rows_per_doc - 1) // max_rows_per_doc

                    lines: list[str] = []
                    for row in chunk:
                        row_text = " | ".join(
                            f"{col}: {row.get(col, '')}" for col in columns if row.get(col)
                        )
                        if title_column and row.get(title_column):
                            lines.append(f"## {row[title_column]}\n{row_text}")
                        else:
                            lines.append(row_text)

                    title = csv_path.stem
                    if total_chunks > 1:
                        title = f"{title} (Part {chunk_num}/{total_chunks})"

                    documents.append(
                        Document(
                            title=title,
                            content="\n\n".join(lines),
                            source=str(csv_path),
                            metadata={
                                "format": "csv",
                                "rows": len(chunk),
                                "chunk": chunk_num,
                                "total_chunks": total_chunks,
                            },
                        )
                    )
            except Exception:
                logger.exception("Failed to load CSV %s", csv_path)

        return documents

    def load_docx(self, directory: str | Path) -> list[Document]:
        """Load Word documents using python-docx."""
        return self.load_directory(directory, formats=["docx"])

    def load_markdown(self, directory: str | Path) -> list[Document]:
        """Load Markdown files directly (GraphRAG native support)."""
        return self.load_directory(directory, formats=["md"])

    def scan_github(
        self,
        repo_url: str,
        patterns: list[str] | None = None,
        branch: str = "main",
        depth: int = 1,
    ) -> list[Document]:
        """Clone a GitHub repository and extract documentation.

        Args:
            repo_url: GitHub repository URL.
            patterns: File patterns to extract. Default: ["*.md", "*.txt", "*.rst"].
            branch: Branch to clone.
            depth: Git clone depth (1 for shallow).

        Returns:
            List of documents extracted from the repository.
        """
        patterns = patterns or ["*.md", "*.txt", "*.rst"]

        with tempfile.TemporaryDirectory(prefix="graphrag-github-") as tmp_dir:
            logger.info("Cloning %s (branch=%s, depth=%d)", repo_url, branch, depth)
            try:
                subprocess.run(
                    [
                        "git",
                        "clone",
                        "--depth",
                        str(depth),
                        "--branch",
                        branch,
                        repo_url,
                        tmp_dir,
                    ],
                    check=True,
                    capture_output=True,
                    text=True,
                    timeout=120,
                )
            except subprocess.CalledProcessError as e:
                logger.error("Git clone failed: %s", e.stderr)
                raise RuntimeError(f"Failed to clone {repo_url}: {e.stderr}") from e

            documents: list[Document] = []
            for pattern in patterns:
                for file_path in glob.glob(
                    os.path.join(tmp_dir, "**", pattern), recursive=True
                ):
                    # Skip .git directory
                    if ".git" in Path(file_path).parts:
                        continue
                    try:
                        content = Path(file_path).read_text(encoding="utf-8", errors="replace")
                        if not content.strip():
                            continue

                        rel_path = os.path.relpath(file_path, tmp_dir)
                        documents.append(
                            Document(
                                title=rel_path,
                                content=content,
                                source=f"{repo_url}/blob/{branch}/{rel_path}",
                                metadata={
                                    "format": Path(file_path).suffix.lstrip("."),
                                    "repo": repo_url,
                                    "branch": branch,
                                    "path": rel_path,
                                },
                            )
                        )
                    except Exception:
                        logger.exception("Failed to read %s", file_path)

            logger.info("Extracted %d documents from %s", len(documents), repo_url)
            return documents

    def upload_to_blob(
        self,
        documents: list[Document],
        container_name: str,
        storage_account_url: str | None = None,
        overwrite: bool = True,
    ) -> dict[str, Any]:
        """Upload processed documents to Azure Blob Storage for GraphRAG.

        GraphRAG expects plain text files in a flat directory structure.

        Args:
            documents: Processed documents to upload.
            container_name: Blob container name (e.g., 'graphrag-input').
            storage_account_url: Storage account URL. Defaults to
                AZURE_STORAGE_ACCOUNT_URL environment variable.
            overwrite: Whether to overwrite existing blobs.

        Returns:
            Upload manifest with file count and total size.
        """
        account_url = storage_account_url or os.getenv("AZURE_STORAGE_ACCOUNT_URL")
        if not account_url:
            raise ValueError(
                "storage_account_url or AZURE_STORAGE_ACCOUNT_URL required"
            )

        blob_service = BlobServiceClient(
            account_url=account_url, credential=self._credential
        )

        # Create container if needed
        try:
            blob_service.create_container(container_name)
            logger.info("Created container: %s", container_name)
        except Exception:
            pass  # Container may already exist

        container_client = blob_service.get_container_client(container_name)
        manifest = {"files": [], "total_size": 0, "count": 0}

        for i, doc in enumerate(documents):
            # GraphRAG wants flat .txt files
            safe_name = doc.title.replace("/", "_").replace("\\", "_")
            safe_name = "".join(c for c in safe_name if c.isalnum() or c in "-_. ")
            blob_name = f"{i:04d}_{safe_name}.txt"

            content = doc.to_text()
            content_bytes = content.encode("utf-8")

            blob_client = container_client.get_blob_client(blob_name)
            blob_client.upload_blob(content_bytes, overwrite=overwrite)

            manifest["files"].append(
                {
                    "blob_name": blob_name,
                    "source": doc.source,
                    "size": len(content_bytes),
                    "title": doc.title,
                }
            )
            manifest["total_size"] += len(content_bytes)
            manifest["count"] += 1

        # Upload manifest
        manifest_blob = container_client.get_blob_client("_manifest.json")
        manifest_blob.upload_blob(
            json.dumps(manifest, indent=2).encode("utf-8"), overwrite=True
        )

        logger.info(
            "Uploaded %d documents (%d bytes) to %s/%s",
            manifest["count"],
            manifest["total_size"],
            account_url,
            container_name,
        )
        return manifest

    def save_to_directory(
        self, documents: list[Document], output_dir: str | Path
    ) -> dict[str, Any]:
        """Save processed documents to a local directory for GraphRAG.

        Alternative to upload_to_blob for local development.

        Args:
            documents: Processed documents.
            output_dir: Directory to write text files to.

        Returns:
            Manifest with file count and paths.
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)

        manifest = {"files": [], "total_size": 0, "count": 0}

        for i, doc in enumerate(documents):
            safe_name = doc.title.replace("/", "_").replace("\\", "_")
            safe_name = "".join(c for c in safe_name if c.isalnum() or c in "-_. ")
            file_name = f"{i:04d}_{safe_name}.txt"
            file_path = output_dir / file_name

            content = doc.to_text()
            file_path.write_text(content, encoding="utf-8")

            manifest["files"].append(
                {
                    "file_name": file_name,
                    "source": doc.source,
                    "size": len(content.encode("utf-8")),
                    "title": doc.title,
                }
            )
            manifest["total_size"] += len(content.encode("utf-8"))
            manifest["count"] += 1

        # Write manifest
        manifest_path = output_dir / "_manifest.json"
        manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")

        logger.info("Saved %d documents to %s", manifest["count"], output_dir)
        return manifest

    # --- Private methods ---

    def _load_file(self, file_path: Path, fmt: str) -> list[Document]:
        """Load a single file based on format."""
        loaders = {
            "pdf": self._load_pdf,
            "csv": self._load_csv_single,
            "docx": self._load_docx_single,
            "md": self._load_text,
            "txt": self._load_text,
            "json": self._load_json,
        }
        loader = loaders.get(fmt)
        if not loader:
            raise ValueError(f"No loader for format: {fmt}")
        return loader(file_path)

    def _load_pdf(self, file_path: Path) -> list[Document]:
        """Load PDF using Azure Document Intelligence."""
        if not self._di_endpoint:
            # Fallback: try PyPDF2
            return self._load_pdf_fallback(file_path)

        from azure.ai.documentintelligence import DocumentIntelligenceClient
        from azure.ai.documentintelligence.models import AnalyzeDocumentRequest

        client = DocumentIntelligenceClient(
            endpoint=self._di_endpoint, credential=self._credential
        )

        with open(file_path, "rb") as f:
            poller = client.begin_analyze_document(
                "prebuilt-read",
                analyze_request=AnalyzeDocumentRequest(bytes_source=f.read()),
            )
            result = poller.result()

        content = result.content if result.content else ""
        return [
            Document(
                title=file_path.stem,
                content=content,
                source=str(file_path),
                metadata={
                    "format": "pdf",
                    "pages": len(result.pages) if result.pages else 0,
                },
            )
        ]

    def _load_pdf_fallback(self, file_path: Path) -> list[Document]:
        """Load PDF using PyPDF2 as fallback."""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            logger.warning(
                "Neither Azure Document Intelligence nor PyPDF2 available. "
                "Install PyPDF2: pip install PyPDF2"
            )
            return []

        reader = PdfReader(str(file_path))
        pages: list[str] = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)

        return [
            Document(
                title=file_path.stem,
                content="\n\n".join(pages),
                source=str(file_path),
                metadata={"format": "pdf", "pages": len(pages)},
            )
        ]

    def _load_csv_single(self, file_path: Path) -> list[Document]:
        """Load a single CSV file as a document."""
        content = file_path.read_text(encoding="utf-8-sig", errors="replace")
        return [
            Document(
                title=file_path.stem,
                content=content,
                source=str(file_path),
                metadata={"format": "csv"},
            )
        ]

    def _load_docx_single(self, file_path: Path) -> list[Document]:
        """Load a Word document using python-docx."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            logger.warning("python-docx not installed. Install: pip install python-docx")
            return []

        doc = DocxDocument(str(file_path))
        paragraphs: list[str] = []
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style and para.style.name.startswith("Heading"):
                    level = para.style.name.replace("Heading ", "")
                    try:
                        hashes = "#" * int(level)
                    except ValueError:
                        hashes = "#"
                    paragraphs.append(f"{hashes} {para.text}")
                else:
                    paragraphs.append(para.text)

        return [
            Document(
                title=file_path.stem,
                content="\n\n".join(paragraphs),
                source=str(file_path),
                metadata={"format": "docx"},
            )
        ]

    def _load_text(self, file_path: Path) -> list[Document]:
        """Load a plain text or markdown file."""
        content = file_path.read_text(encoding="utf-8", errors="replace")
        return [
            Document(
                title=file_path.stem,
                content=content,
                source=str(file_path),
                metadata={"format": file_path.suffix.lstrip(".")},
            )
        ]

    def _load_json(self, file_path: Path) -> list[Document]:
        """Load a JSON file, extracting text content."""
        raw = file_path.read_text(encoding="utf-8", errors="replace")
        try:
            data = json.loads(raw)
        except json.JSONDecodeError:
            logger.warning("Invalid JSON: %s", file_path)
            return []

        if isinstance(data, list):
            # Array of objects: each becomes a document
            documents: list[Document] = []
            for i, item in enumerate(data):
                if isinstance(item, dict):
                    title = item.get("title", item.get("name", f"{file_path.stem}_{i}"))
                    content = item.get(
                        "content",
                        item.get("text", json.dumps(item, indent=2)),
                    )
                    documents.append(
                        Document(
                            title=str(title),
                            content=str(content),
                            source=str(file_path),
                            metadata={"format": "json", "index": i},
                        )
                    )
            return documents
        elif isinstance(data, dict):
            title = data.get("title", data.get("name", file_path.stem))
            content = data.get("content", data.get("text", json.dumps(data, indent=2)))
            return [
                Document(
                    title=str(title),
                    content=str(content),
                    source=str(file_path),
                    metadata={"format": "json"},
                )
            ]
        return []
